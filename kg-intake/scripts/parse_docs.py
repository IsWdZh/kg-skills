# -*- coding: utf-8 -*-
"""文档接入解析脚本：HTML/docx/PDF -> 带段落锚点的结构化 Markdown + 文档台账。

用法: python parse_docs.py <源文档根目录> <工作区目录>
源目录下按来源分子文件夹；本脚本内置 来源文件夹名->来源码/地区 的映射，
迁移到其他业务时只需修改 SOURCE_MAP。
"""
import sys, re, html, hashlib, io
from datetime import date
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")

SOURCE_MAP = {  # 文件夹名: (来源码, 地区, 发布机构)
    "省公积金": ("PROV", "甘肃省直", "甘肃省住房资金管理中心"),
    "兰州": ("LZ", "兰州市", "兰州住房公积金管理中心"),
    "嘉峪关": ("JYG", "嘉峪关市", "嘉峪关市住房公积金管理中心"),
}

DOC_TYPE_RULES = [  # (关键词正则, 文档类型)
    (r"办事指南|须知|须准备", "办事指南"),
    (r"操作规程|办理规范", "业务规程"),
    (r"通知|办法|规定|文件|印发", "政策文件"),
    (r"您知道吗|可以这样", "宣传问答"),
]


def detect_doc_type(title: str) -> str:
    for pat, t in DOC_TYPE_RULES:
        if re.search(pat, title):
            return t
    return "其他"


def detect_publish_date(text: str) -> str:
    m = re.search(r"发布日期[：:]\s*(\d{4}-\d{1,2}-\d{1,2})", text)
    if m:
        return m.group(1)
    m = re.search(r"(\d{4})年(\d{1,2})月(\d{1,2})日", text[-600:])  # 落款日期常在结尾
    if m:
        return f"{m.group(1)}-{int(m.group(2)):02d}-{int(m.group(3)):02d}"
    m = re.search(r"[〔\[](\d{4})[〕\]]\s*\d+\s*号", text)
    if m:
        return f"{m.group(1)}(文号年份)"
    return "未知"


def extract_html(path: Path) -> list[str]:
    from bs4 import BeautifulSoup, Comment
    raw = path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(raw, "lxml")
    for tag in soup(["script", "style", "head"]):
        tag.decompose()
    for c in soup.find_all(string=lambda s: isinstance(s, Comment)):
        c.extract()
    blocks = []
    # 政务页面常见正文容器优先；找不到再用 body
    body = (soup.select_one("div#zoom") or soup.select_one("div.article-conter")
            or soup.select_one("div.bt_content") or soup.body or soup)
    img_count = len(body.find_all("img"))
    if img_count:
        blocks.append(f"〔注：本页含 {img_count} 张图片，文字可能在图片中，需视觉识别〕")
    for el in body.find_all(["h1", "h2", "h3", "h4", "p", "li", "table"]):
        if el.attrs.get("__consumed__"):
            continue
        if el.name == "table":
            # 含多个 <p> 的视为布局表格，跳过让内部段落正常输出
            if len(el.find_all("p")) > 3:
                continue
            rows = []
            for tr in el.find_all("tr"):
                cells = [td.get_text(" ", strip=True) for td in tr.find_all(["td", "th"])]
                if any(cells):
                    rows.append("| " + " | ".join(cells) + " |")
            if rows:
                header_sep = "|" + "---|" * (rows[0].count("|") - 1)
                blocks.append("\n".join([rows[0], header_sep] + rows[1:]))
            for sub in el.find_all(["p", "li"]):
                sub.attrs["__consumed__"] = "1"
            continue
        if el.find("table"):
            continue
        txt = el.get_text(" ", strip=True)
        if txt:
            blocks.append(html.unescape(txt))
    if not blocks:  # 无块级结构时退化为按行切
        text = re.sub(r"<[^>]+>", "\n", raw)
        blocks = [l.strip() for l in html.unescape(text).split("\n") if l.strip()]
    return blocks


def extract_docx(path: Path) -> list[str]:
    import docx
    d = docx.Document(str(path))
    blocks = []
    for para in d.paragraphs:
        t = para.text.strip()
        if t:
            blocks.append(t)
    for tb in d.tables:
        rows = []
        for r in tb.rows:
            cells = [c.text.strip().replace("\n", " ") for c in r.cells]
            if any(cells):
                rows.append("| " + " | ".join(cells) + " |")
        if rows:
            header_sep = "|" + "---|" * (rows[0].count("|") - 1)
            blocks.append("\n".join([rows[0], header_sep] + rows[1:]))
    return blocks


def extract_pdf(path: Path) -> list[str]:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    blocks = []
    for i, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        paras, buf = [], []
        for line in text.split("\n"):
            line = line.strip()
            if not line:
                if buf:
                    paras.append(" ".join(buf)); buf = []
                continue
            buf.append(line)
            if re.search(r"[。；：]$", line):  # 句末断段
                paras.append(" ".join(buf)); buf = []
        if buf:
            paras.append(" ".join(buf))
        for p in paras:
            if p.strip():
                blocks.append(f"〔页{i}〕{p.strip()}")
    return blocks


def parse_front_matter(text: str):
    m = re.match(r"^---\n(.*?)\n---\n?", text, re.S)
    meta, body = {}, text
    if m:
        for line in m.group(1).split("\n"):
            if ":" in line:
                k, v = line.split(":", 1)
                meta[k.strip()] = v.strip()
        body = text[m.end():]
    return meta, body


def _write_registry(ws: Path, registry):
    reg_dir = ws / "00-文档台账"
    reg_dir.mkdir(parents=True, exist_ok=True)
    out = io.StringIO()
    out.write("# 文档台账\n\n")
    out.write(f"共登记 {len(registry)} 份文档。状态说明：`已接入`=登记并解析完成，待打标。\n\n")
    out.write("| 编号 | 标题 | 适用地区 | 文档类型 | 发布时间 | 格式 | 段落数 | 字数 | 完整性 | 去重 | 源文件 |\n")
    out.write("|---|---|---|---|---|---|---|---|---|---|---|\n")
    for r in registry:
        out.write("| " + " | ".join(str(x) for x in r) + " |\n")
    (reg_dir / "台账.md").write_text(out.getvalue(), encoding="utf-8")
    print(f"\n台账已写入 {reg_dir / '台账.md'}")


def main_from_md(ws: Path):
    """从 00-原始MD（kg-normalize 归一产物）解析为带锚点结构化文档 + 台账。"""
    registry, hash_seen = [], {}
    batch_id = f"INTAKE-{date.today().isoformat()}"
    src_root = ws / "00-原始MD"
    if not src_root.exists():
        print("未找到 00-原始MD/，请先运行 to_markdown.py 完成格式归一（kg-normalize 环节）")
        return
    parsed_root = ws / "10-解析文档"
    for code_dir in sorted(p for p in src_root.iterdir() if p.is_dir() and p.name != "_图片"):
        code = code_dir.name
        out_dir = parsed_root / code
        out_dir.mkdir(parents=True, exist_ok=True)
        for f in sorted(code_dir.glob("*.md")):
            meta, body = parse_front_matter(f.read_text(encoding="utf-8"))
            doc_id = meta.get("编号", f.stem)
            title = meta.get("标题", doc_id)
            pub = meta.get("发布时间", "未知")
            未转写 = "待视觉转写" in meta.get("归一状态", "")
            paras = []
            for line in body.split("\n"):
                s = line.strip()
                if not s or s.startswith("# "):
                    continue
                paras.append(re.sub(r"^>\s*", "", s))
            full = "".join(paras)
            h = hashlib.sha256(re.sub(r"\s+", "", full).encode()).hexdigest()
            dup = hash_seen.get(h, "")
            if not dup:
                hash_seen[h] = doc_id
            dtype = detect_doc_type(title)
            chars = len(re.sub(r"\s+", "", full))
            comp = []
            if 未转写:
                comp.append("图片型-未转写（需回 kg-normalize 补全）")
            else:
                if pub == "未知":
                    comp.append("缺发布时间")
                if chars < 100:
                    comp.append("正文过短")
            comp = "；".join(comp) if comp else "完整"
            lines = ["---", f"编号: {doc_id}", f"标题: {title}",
                     f"地区: {meta.get('地区','')}", f"发布机构: {meta.get('发布机构','')}",
                     f"文档类型: {dtype}", f"发布时间: {pub}", "生效时间: 未知",
                     f"源文件: {meta.get('源文件','')}", f"内容哈希: {h}",
                     "内容哈希算法: sha256-normalized-v1", f"接入批次: {batch_id}",
                     f"接入时间: {date.today().isoformat()}", "来源URL: 未知",
                     "权限范围: 公开", "责任人: 未指定", f"完整性: {comp}",
                     "上游: 00-原始MD（格式归一产物）", "---", "", f"# {title}", ""]
            for i, p in enumerate(paras, 1):
                lines.append(f"[p{i:03d}] {p}")
            (out_dir / f"{doc_id}.md").write_text("\n".join(lines), encoding="utf-8")
            registry.append((doc_id, title, meta.get("地区", ""), dtype, pub,
                             meta.get("原始格式", ""), len(paras), chars, comp,
                             f"重复于{dup}" if dup else "唯一", meta.get("源文件", "")))
            print(f"{doc_id}  {comp:<16} {len(paras):>3}段 {chars:>6}字  {title[:26]}")
    _write_registry(ws, registry)


def main(src_root: Path, ws: Path):
    registry, hash_seen = [], {}
    batch_id = f"INTAKE-{date.today().isoformat()}"
    parsed_root = ws / "10-解析文档"
    for folder, (code, region, org) in SOURCE_MAP.items():
        d = src_root / folder
        if not d.exists():
            continue
        out_dir = parsed_root / code
        out_dir.mkdir(parents=True, exist_ok=True)
        files = sorted(d.iterdir(), key=lambda p: p.name)
        for idx, f in enumerate([x for x in files if x.suffix.lower() in (".html", ".htm", ".docx", ".pdf")], 1):
            doc_id = f"{code}-{idx:03d}"
            title = re.sub(r"^\d+\.", "", f.stem).strip()
            try:
                if f.suffix.lower() in (".html", ".htm"):
                    blocks = extract_html(f)
                elif f.suffix.lower() == ".docx":
                    blocks = extract_docx(f)
                else:
                    blocks = extract_pdf(f)
                err = ""
            except Exception as e:
                blocks, err = [], f"{type(e).__name__}: {e}"
            full = "\n".join(blocks)
            h = hashlib.sha256(re.sub(r"\s+", "", full).encode()).hexdigest()
            source_sha256 = hashlib.sha256(f.read_bytes()).hexdigest()
            dup = hash_seen.get(h, "")
            if not dup:
                hash_seen[h] = doc_id
            pub = detect_publish_date(full)
            if pub == "未知" and f.suffix.lower() in (".html", ".htm"):
                # 发布日期常在正文容器之外的页头区域
                raw_text = re.sub(r"<[^>]+>", " ", f.read_text(encoding="utf-8", errors="ignore")[:20000])
                pub = detect_publish_date(raw_text)
            dtype = detect_doc_type(title)
            chars = len(full)
            completeness = []
            if pub == "未知":
                completeness.append("缺发布时间")
            if chars < 100:
                completeness.append("正文过短")
            if err:
                completeness.append("解析失败")
            comp = "；".join(completeness) if completeness else "完整"
            lines = [
                "---",
                f"编号: {doc_id}",
                f"标题: {title}",
                f"地区: {region}",
                f"发布机构: {org}",
                f"文档类型: {dtype}",
                f"发布时间: {pub}",
                "生效时间: 未知",
                f"源文件: {f.name}",
                f"源文件SHA256: {source_sha256}",
                f"内容哈希: {h}",
                "内容哈希算法: sha256-normalized-v1",
                f"接入批次: {batch_id}",
                f"接入时间: {date.today().isoformat()}",
                "来源URL: 未知",
                "权限范围: 公开",
                "责任人: 未指定",
                f"完整性: {comp}",
                "---",
                "",
                f"# {title}",
                "",
            ]
            for i, b in enumerate(blocks, 1):
                if "\n" in b:  # 表格块
                    lines.append(f"[p{i:03d}]\n{b}\n")
                else:
                    lines.append(f"[p{i:03d}] {b}")
            (out_dir / f"{doc_id}.md").write_text("\n".join(lines), encoding="utf-8")
            registry.append((doc_id, title, region, dtype, pub, f.suffix.lower()[1:],
                             len(blocks), chars, comp, f"重复于{dup}" if dup else "唯一", f.name))
            print(f"{doc_id}  {comp:<6} {len(blocks):>3}段 {chars:>6}字  {title[:30]}")

    _write_registry(ws, registry)


if __name__ == "__main__":
    # 推荐流程：先 kg-normalize 出 00-原始MD，再 --from-md 做结构化解析
    #   python parse_docs.py --from-md <工作区目录>
    # 兼容旧流程（直接从原始文件解析）：
    #   python parse_docs.py <源文档根目录> <工作区目录>
    if "--from-md" in sys.argv:
        pos = [a for a in sys.argv[1:] if not a.startswith("--")]
        main_from_md(Path(pos[0]))
    else:
        main(Path(sys.argv[1]), Path(sys.argv[2]))
